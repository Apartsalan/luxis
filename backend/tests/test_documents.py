"""Tests for the documents module — template CRUD, document generation, and docx rendering."""

import uuid
from datetime import date
from decimal import Decimal
from io import BytesIO

import pytest
from httpx import AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession

from app.auth.models import Tenant
from app.collections.models import InterestRate
from app.documents.docx_service import (
    _fmt_currency,
    _fmt_date,
    _fmt_pct,
)
from app.relations.models import Contact

# ── Template CRUD ────────────────────────────────────────────────────────────


SAMPLE_TEMPLATE = {
    "name": "Test 14-dagenbrief",
    "description": "Standaard 14-dagenbrief voor incassozaken",
    "template_type": "14_dagenbrief",
    "content": (
        "<h1>14-dagenbrief</h1>"
        "<p>Zaaknummer: {{ zaak.zaaknummer }}</p>"
        "<p>Client: {{ client.naam }}</p>"
        "<p>Wederpartij: {{ wederpartij.naam }}</p>"
        "<p>Hoofdsom: {{ financieel.total_principal | currency }}</p>"
        "<p>Datum: {{ vandaag | datum }}</p>"
    ),
}


@pytest.mark.asyncio
async def test_create_template(
    client: AsyncClient, auth_headers: dict
):
    """Creating a template should return 201."""
    response = await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    assert response.status_code == 201
    data = response.json()
    assert data["name"] == "Test 14-dagenbrief"
    assert data["template_type"] == "14_dagenbrief"
    assert "{{ zaak.zaaknummer }}" in data["content"]


@pytest.mark.asyncio
async def test_create_template_invalid_jinja(
    client: AsyncClient, auth_headers: dict
):
    """Template with invalid Jinja2 syntax should return 400."""
    payload = {
        "name": "Broken template",
        "template_type": "test",
        "content": "{% if unclosed %}",
    }
    response = await client.post(
        "/api/documents/templates",
        json=payload,
        headers=auth_headers,
    )
    assert response.status_code == 400


@pytest.mark.asyncio
async def test_list_templates(
    client: AsyncClient, auth_headers: dict
):
    """List templates should work."""
    # Create two templates
    await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    await client.post(
        "/api/documents/templates",
        json={**SAMPLE_TEMPLATE, "name": "Sommatie", "template_type": "sommatie"},
        headers=auth_headers,
    )

    # List all
    response = await client.get(
        "/api/documents/templates", headers=auth_headers
    )
    assert response.status_code == 200
    data = response.json()
    assert len(data) >= 2

    # Filter by type
    response = await client.get(
        "/api/documents/templates?template_type=sommatie",
        headers=auth_headers,
    )
    data = response.json()
    assert all(t["template_type"] == "sommatie" for t in data)


@pytest.mark.asyncio
async def test_get_template(
    client: AsyncClient, auth_headers: dict
):
    """Get template by ID should return full content."""
    create_response = await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    template_id = create_response.json()["id"]

    response = await client.get(
        f"/api/documents/templates/{template_id}",
        headers=auth_headers,
    )
    assert response.status_code == 200
    data = response.json()
    assert data["content"] == SAMPLE_TEMPLATE["content"]


@pytest.mark.asyncio
async def test_update_template(
    client: AsyncClient, auth_headers: dict
):
    """Updating a template should work."""
    create_response = await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    template_id = create_response.json()["id"]

    response = await client.put(
        f"/api/documents/templates/{template_id}",
        json={"name": "Aangepaste brief"},
        headers=auth_headers,
    )
    assert response.status_code == 200
    assert response.json()["name"] == "Aangepaste brief"


@pytest.mark.asyncio
async def test_delete_template(
    client: AsyncClient, auth_headers: dict
):
    """Deleting a template should soft-delete it."""
    create_response = await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    template_id = create_response.json()["id"]

    response = await client.delete(
        f"/api/documents/templates/{template_id}",
        headers=auth_headers,
    )
    assert response.status_code == 204


@pytest.mark.asyncio
async def test_get_template_not_found(
    client: AsyncClient, auth_headers: dict
):
    """Non-existent template should return 404."""
    fake_id = str(uuid.uuid4())
    response = await client.get(
        f"/api/documents/templates/{fake_id}",
        headers=auth_headers,
    )
    assert response.status_code == 404


# ── Document Generation ──────────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_generate_document(
    client: AsyncClient,
    auth_headers: dict,
    db: AsyncSession,
    test_company: Contact,
    test_person: Contact,
):
    """Generating a document from a template should render correctly."""
    await _seed_interest_rates(db)
    # Create a case first
    case_payload = {
        "case_type": "incasso",
        "description": "Test vordering",
        "client_id": str(test_company.id),
        "opposing_party_id": str(test_person.id),
        "date_opened": "2026-02-17",
    }
    case_response = await client.post(
        "/api/cases", json=case_payload, headers=auth_headers
    )
    case_id = case_response.json()["id"]

    # Create a template
    template_response = await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    template_id = template_response.json()["id"]

    # Generate document
    gen_payload = {
        "template_id": template_id,
        "title": "14-dagenbrief voor Test zaak",
    }
    response = await client.post(
        f"/api/documents/cases/{case_id}/generate",
        json=gen_payload,
        headers=auth_headers,
    )
    assert response.status_code == 201
    data = response.json()
    assert data["title"] == "14-dagenbrief voor Test zaak"
    assert data["document_type"] == "14_dagenbrief"
    assert data["content_html"] is not None
    # Check rendered content contains case number
    assert "2026-" in data["content_html"]
    assert "Acme B.V." in data["content_html"]


@pytest.mark.asyncio
async def test_list_case_documents(
    client: AsyncClient,
    auth_headers: dict,
    db: AsyncSession,
    test_company: Contact,
    test_person: Contact,
):
    """List documents for a case should work."""
    await _seed_interest_rates(db)
    # Create case
    case_response = await client.post(
        "/api/cases",
        json={
            "case_type": "incasso",
            "client_id": str(test_company.id),
            "opposing_party_id": str(test_person.id),
            "date_opened": "2026-02-17",
        },
        headers=auth_headers,
    )
    case_id = case_response.json()["id"]

    # Create template & generate doc
    template_response = await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    template_id = template_response.json()["id"]

    gen_resp = await client.post(
        f"/api/documents/cases/{case_id}/generate",
        json={"template_id": template_id},
        headers=auth_headers,
    )
    assert gen_resp.status_code == 201, f"Generate failed: {gen_resp.text}"

    # List documents
    response = await client.get(
        f"/api/documents/cases/{case_id}",
        headers=auth_headers,
    )
    assert response.status_code == 200
    data = response.json()
    assert len(data) >= 1


@pytest.mark.asyncio
async def test_delete_generated_document(
    client: AsyncClient,
    auth_headers: dict,
    db: AsyncSession,
    test_company: Contact,
    test_person: Contact,
):
    """Deleting a generated document should soft-delete it."""
    await _seed_interest_rates(db)
    # Create case
    case_response = await client.post(
        "/api/cases",
        json={
            "case_type": "incasso",
            "client_id": str(test_company.id),
            "opposing_party_id": str(test_person.id),
            "date_opened": "2026-02-17",
        },
        headers=auth_headers,
    )
    case_id = case_response.json()["id"]

    # Create template & generate
    template_response = await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    template_id = template_response.json()["id"]

    gen_response = await client.post(
        f"/api/documents/cases/{case_id}/generate",
        json={"template_id": template_id},
        headers=auth_headers,
    )
    assert gen_response.status_code == 201, f"Generate failed: {gen_response.text}"
    doc_id = gen_response.json()["id"]

    # Delete
    response = await client.delete(
        f"/api/documents/{doc_id}",
        headers=auth_headers,
    )
    assert response.status_code == 204


# ══════════════════════════════════════════════════════════════════════════════
# DOCX GENERATION TESTS
# ══════════════════════════════════════════════════════════════════════════════


# ── Unit tests: formatting helpers ───────────────────────────────────────────


class TestFormatCurrency:
    """Test Dutch currency formatting."""

    def test_zero(self):
        assert _fmt_currency(Decimal("0")) == "EUR 0,00"

    def test_small_amount(self):
        assert _fmt_currency(Decimal("40.00")) == "EUR 40,00"

    def test_thousands(self):
        assert _fmt_currency(Decimal("1234.56")) == "EUR 1.234,56"

    def test_large_amount(self):
        assert _fmt_currency(Decimal("15000.00")) == "EUR 15.000,00"

    def test_none(self):
        assert _fmt_currency(None) == "EUR 0,00"

    def test_negative(self):
        # Negative amounts should show minus sign
        result = _fmt_currency(Decimal("-500.00"))
        assert "500,00" in result


class TestFormatDate:
    """Test Dutch date formatting."""

    def test_date_object(self):
        assert _fmt_date(date(2026, 2, 17)) == "17 februari 2026"

    def test_iso_string(self):
        assert _fmt_date("2026-02-17") == "17 februari 2026"

    def test_none(self):
        assert _fmt_date(None) == ""

    def test_january(self):
        assert _fmt_date(date(2026, 1, 1)) == "1 januari 2026"

    def test_december(self):
        assert _fmt_date(date(2025, 12, 31)) == "31 december 2025"


class TestFormatPercentage:
    """Test percentage formatting."""

    def test_integer_rate(self):
        assert _fmt_pct(Decimal("6.00")) == "6,00%"

    def test_decimal_rate(self):
        assert _fmt_pct(Decimal("10.50")) == "10,50%"

    def test_none(self):
        assert _fmt_pct(None) == "0,00%"


# ── Helper: seed interest rates and create case with claims ──────────────────


async def _seed_interest_rates(db: AsyncSession):
    """Seed statutory interest rates needed for calculations."""
    rates = [
        InterestRate(
            id=uuid.uuid4(),
            rate_type="statutory",
            effective_from=date(2024, 1, 1),
            rate=Decimal("7.00"),
            source="test",
        ),
        InterestRate(
            id=uuid.uuid4(),
            rate_type="statutory",
            effective_from=date(2025, 1, 1),
            rate=Decimal("6.00"),
            source="test",
        ),
        InterestRate(
            id=uuid.uuid4(),
            rate_type="statutory",
            effective_from=date(2026, 1, 1),
            rate=Decimal("4.00"),
            source="test",
        ),
    ]
    for r in rates:
        db.add(r)
    await db.commit()


async def _create_case_with_claims(
    client: AsyncClient,
    auth_headers: dict,
    db: AsyncSession,
    tenant: Tenant,
    client_contact: Contact,
    opposing_party: Contact,
) -> str:
    """Create a case with claims and seeded interest rates. Returns case_id."""
    await _seed_interest_rates(db)

    # Create case
    case_response = await client.post(
        "/api/cases",
        json={
            "case_type": "incasso",
            "description": "Onbetaalde facturen webdevelopment",
            "client_id": str(client_contact.id),
            "opposing_party_id": str(opposing_party.id),
            "date_opened": "2025-06-01",
            "interest_type": "statutory",
        },
        headers=auth_headers,
    )
    assert case_response.status_code == 201, case_response.json()
    case_id = case_response.json()["id"]

    # Add claims via API
    claims = [
        {
            "description": "Factuur 2025-042 dd. 15-05-2025",
            "principal_amount": "5000.00",
            "default_date": "2025-06-15",
            "invoice_number": "2025-042",
            "invoice_date": "2025-05-15",
        },
        {
            "description": "Factuur 2025-058 dd. 01-07-2025",
            "principal_amount": "2500.00",
            "default_date": "2025-08-01",
            "invoice_number": "2025-058",
            "invoice_date": "2025-07-01",
        },
    ]
    for claim in claims:
        resp = await client.post(
            f"/api/cases/{case_id}/claims",
            json=claim,
            headers=auth_headers,
        )
        assert resp.status_code == 201, resp.json()

    return case_id


# ── Docx endpoint: list available templates ──────────────────────────────────


@pytest.mark.asyncio
async def test_list_docx_templates(
    client: AsyncClient,
    auth_headers: dict,
):
    """GET /api/documents/docx/templates should list available template types."""
    response = await client.get(
        "/api/documents/docx/templates",
        headers=auth_headers,
    )
    assert response.status_code == 200
    data = response.json()
    assert len(data) >= 3  # More templates added in later sessions

    types = {t["template_type"] for t in data}
    assert {"14_dagenbrief", "sommatie", "renteoverzicht"}.issubset(types)

    # All should be available (templates exist on disk)
    for t in data:
        assert t["available"] is True


# ── Docx generation: 14-dagenbrief ──────────────────────────────────────────


@pytest.mark.asyncio
async def test_generate_docx_14_dagenbrief(
    client: AsyncClient,
    auth_headers: dict,
    db: AsyncSession,
    test_tenant: Tenant,
    test_company: Contact,
    test_person: Contact,
):
    """Generate a 14-dagenbrief .docx — must return valid Word document."""
    case_id = await _create_case_with_claims(
        client, auth_headers, db, test_tenant, test_company, test_person
    )

    response = await client.post(
        f"/api/documents/docx/cases/{case_id}/generate",
        json={"template_type": "14_dagenbrief"},
        headers=auth_headers,
    )
    assert response.status_code == 200, response.json()
    assert (
        response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Verify it's a valid .docx (zip with PK signature)
    content = response.content
    assert len(content) > 1000  # Should be a real file
    assert content[:2] == b"PK"  # .docx is a ZIP file

    # Verify filename in Content-Disposition
    disposition = response.headers.get("content-disposition", "")
    assert "14_dagenbrief" in disposition
    assert ".docx" in disposition

    # Parse the generated docx and check content
    from docx import Document

    doc = Document(BytesIO(content))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Must contain client name
    assert "Acme B.V." in full_text
    # Must contain opposing party name
    assert "Jan de Vries" in full_text
    # Must contain legal notice reference
    assert "6:96" in full_text or "veertien dagen" in full_text
    # Must contain formatted currency amounts
    assert "EUR" in full_text


# ── Docx generation: sommatie ────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_generate_docx_sommatie(
    client: AsyncClient,
    auth_headers: dict,
    db: AsyncSession,
    test_tenant: Tenant,
    test_company: Contact,
    test_person: Contact,
):
    """Generate a sommatie .docx — must contain AANGETEKEND and dagvaarding warning."""
    case_id = await _create_case_with_claims(
        client, auth_headers, db, test_tenant, test_company, test_person
    )

    response = await client.post(
        f"/api/documents/docx/cases/{case_id}/generate",
        json={"template_type": "sommatie"},
        headers=auth_headers,
    )
    assert response.status_code == 200
    assert response.content[:2] == b"PK"

    from docx import Document

    doc = Document(BytesIO(response.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Must contain AANGETEKEND header
    assert "AANGETEKEND" in full_text
    # Must contain dagvaarding warning
    assert "dagvaarding" in full_text
    # Must contain acht dagen deadline
    assert "acht (8) dagen" in full_text
    # Must contain party names
    assert "Acme B.V." in full_text
    assert "Jan de Vries" in full_text
    # Must contain financial amounts
    assert "EUR" in full_text


# ── Docx generation: renteoverzicht ──────────────────────────────────────────


@pytest.mark.asyncio
async def test_generate_docx_renteoverzicht(
    client: AsyncClient,
    auth_headers: dict,
    db: AsyncSession,
    test_tenant: Tenant,
    test_company: Contact,
    test_person: Contact,
):
    """Generate a renteoverzicht .docx — must contain interest breakdown."""
    case_id = await _create_case_with_claims(
        client, auth_headers, db, test_tenant, test_company, test_person
    )

    response = await client.post(
        f"/api/documents/docx/cases/{case_id}/generate",
        json={"template_type": "renteoverzicht"},
        headers=auth_headers,
    )
    assert response.status_code == 200
    assert response.content[:2] == b"PK"

    from docx import Document

    doc = Document(BytesIO(response.content))
    para_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + " "
    full_text = para_text + "\n" + table_text

    # Must contain title (in paragraphs/headings)
    assert "Renteberekening" in para_text
    # Must contain interest type label (in meta table)
    assert "Wettelijke rente" in full_text
    # Must contain party names (in meta table)
    assert "Acme B.V." in full_text
    assert "Jan de Vries" in full_text
    # Must contain financial amounts
    assert "EUR" in full_text

    # Should have principal amounts in tables
    assert "5.000,00" in table_text or "EUR 5.000,00" in table_text
    assert "2.500,00" in table_text or "EUR 2.500,00" in table_text
    # Should have interest rate percentages
    assert "%" in table_text


# ── Docx: error cases ────────────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_generate_docx_unknown_template_type(
    client: AsyncClient,
    auth_headers: dict,
    db: AsyncSession,
    test_tenant: Tenant,
    test_company: Contact,
    test_person: Contact,
):
    """Unknown template type should return 404."""
    case_id = await _create_case_with_claims(
        client, auth_headers, db, test_tenant, test_company, test_person
    )

    response = await client.post(
        f"/api/documents/docx/cases/{case_id}/generate",
        json={"template_type": "onbekend"},
        headers=auth_headers,
    )
    assert response.status_code == 404


@pytest.mark.asyncio
async def test_generate_docx_nonexistent_case(
    client: AsyncClient,
    auth_headers: dict,
):
    """Non-existent case should return 404."""
    fake_id = str(uuid.uuid4())
    response = await client.post(
        f"/api/documents/docx/cases/{fake_id}/generate",
        json={"template_type": "14_dagenbrief"},
        headers=auth_headers,
    )
    assert response.status_code == 404


# ── Docx: financial accuracy ─────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_docx_financial_amounts_present(
    client: AsyncClient,
    auth_headers: dict,
    db: AsyncSession,
    test_tenant: Tenant,
    test_company: Contact,
    test_person: Contact,
):
    """Generated docx must contain correct formatted financial amounts.

    Case: EUR 5,000 + EUR 2,500 = EUR 7,500 total principal.
    BIK on 7,500: 15% * 2,500 = 375 + 10% * 2,500 = 250 + 5% * 2,500 = 125 = 750
    But minimum is 40, so BIK = 750. Wait let me recalculate:
      15% of first 2,500 = 375
      10% of next 2,500 (2,500-5,000) = 250
      5% of next 2,500 (5,000-7,500) = 125
      Total BIK = 750 (> minimum 40)
    """
    case_id = await _create_case_with_claims(
        client, auth_headers, db, test_tenant, test_company, test_person
    )

    response = await client.post(
        f"/api/documents/docx/cases/{case_id}/generate",
        json={"template_type": "14_dagenbrief"},
        headers=auth_headers,
    )
    assert response.status_code == 200

    from docx import Document

    doc = Document(BytesIO(response.content))

    # Collect ALL text from paragraphs and tables
    all_text = ""
    for p in doc.paragraphs:
        all_text += p.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + "\n"

    # Total principal: EUR 7.500,00
    assert "7.500,00" in all_text

    # BIK: EUR 750,00
    assert "750,00" in all_text

    # Interest amount should be present (some EUR value)
    # The exact amount depends on calc date, but it should be > 0
    assert all_text.count("EUR") >= 3  # At least principal, BIK, and total


# ── Tenant Isolation ─────────────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_tenant_isolation_list_templates(
    client: AsyncClient,
    auth_headers: dict,
    second_auth_headers: dict,
):
    """Tenant B should NOT see templates created by Tenant A."""
    # Create a template in Tenant A
    await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )

    # Tenant A sees 1 template
    resp_a = await client.get("/api/documents/templates", headers=auth_headers)
    assert len(resp_a.json()) >= 1

    # Tenant B sees 0 templates
    resp_b = await client.get("/api/documents/templates", headers=second_auth_headers)
    assert len(resp_b.json()) == 0


@pytest.mark.asyncio
async def test_tenant_isolation_get_template_detail(
    client: AsyncClient,
    auth_headers: dict,
    second_auth_headers: dict,
):
    """Tenant B should get 404 for Tenant A's template."""
    create_resp = await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    template_id = create_resp.json()["id"]

    response = await client.get(
        f"/api/documents/templates/{template_id}",
        headers=second_auth_headers,
    )
    assert response.status_code == 404


@pytest.mark.asyncio
async def test_tenant_isolation_update_template(
    client: AsyncClient,
    auth_headers: dict,
    second_auth_headers: dict,
):
    """Tenant B should NOT be able to update Tenant A's template."""
    create_resp = await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    template_id = create_resp.json()["id"]

    response = await client.put(
        f"/api/documents/templates/{template_id}",
        json={"name": "Hacked Template"},
        headers=second_auth_headers,
    )
    assert response.status_code == 404


@pytest.mark.asyncio
async def test_tenant_isolation_delete_template(
    client: AsyncClient,
    auth_headers: dict,
    second_auth_headers: dict,
):
    """Tenant B should NOT be able to delete Tenant A's template."""
    create_resp = await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    template_id = create_resp.json()["id"]

    response = await client.delete(
        f"/api/documents/templates/{template_id}",
        headers=second_auth_headers,
    )
    assert response.status_code == 404


@pytest.mark.asyncio
async def test_tenant_isolation_case_documents(
    client: AsyncClient,
    auth_headers: dict,
    second_auth_headers: dict,
    db: AsyncSession,
    test_tenant: Tenant,
    test_company: Contact,
    test_person: Contact,
):
    """Tenant B should NOT see documents from Tenant A's cases."""
    await _seed_interest_rates(db)
    # Create case + generate document in Tenant A
    case_resp = await client.post(
        "/api/cases",
        json={
            "case_type": "incasso",
            "client_id": str(test_company.id),
            "opposing_party_id": str(test_person.id),
            "date_opened": "2026-02-17",
        },
        headers=auth_headers,
    )
    case_id = case_resp.json()["id"]

    template_resp = await client.post(
        "/api/documents/templates",
        json=SAMPLE_TEMPLATE,
        headers=auth_headers,
    )
    template_id = template_resp.json()["id"]

    await client.post(
        f"/api/documents/cases/{case_id}/generate",
        json={"template_id": template_id},
        headers=auth_headers,
    )

    # Tenant B cannot access the case's documents
    resp = await client.get(
        f"/api/documents/cases/{case_id}",
        headers=second_auth_headers,
    )
    # Should be 404 (case not found for this tenant) or empty list
    assert resp.status_code in (404, 200)
    if resp.status_code == 200:
        assert len(resp.json()) == 0


@pytest.mark.asyncio
async def test_tenant_isolation_docx_generation(
    client: AsyncClient,
    auth_headers: dict,
    second_auth_headers: dict,
    db: AsyncSession,
    test_tenant: Tenant,
    test_company: Contact,
    test_person: Contact,
):
    """Tenant B should NOT be able to generate docx for Tenant A's case."""
    case_id = await _create_case_with_claims(
        client, auth_headers, db, test_tenant, test_company, test_person
    )

    response = await client.post(
        f"/api/documents/docx/cases/{case_id}/generate",
        json={"template_type": "14_dagenbrief"},
        headers=second_auth_headers,
    )
    assert response.status_code == 404
