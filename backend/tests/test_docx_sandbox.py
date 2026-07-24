"""SEC-25 wachter: DOCX/HTML-sjablonen worden in een Jinja-sandbox gerenderd.

Bewijst dat een SSTI-payload in een (door gebruiker te uploaden) sjabloon geen
servercode kan bereiken, én dat een gewoon merge-veld nog rendert. Dekt zowel het
DOCX-pad (docx_service) als het oude HTML-pad (service.jinja_env).
"""

import io

import pytest
from docx import Document
from docxtpl import DocxTemplate
from jinja2.exceptions import SecurityError

from app.documents.docx_service import _SANDBOX_ENV
from app.documents.service import jinja_env


def _docx_with(text: str) -> io.BytesIO:
    d = Document()
    d.add_paragraph(text)
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


def test_docx_sandbox_blocks_ssti():
    """De klassieke escape (`__globals__` → builtins → RCE) wordt geblokkeerd."""
    tpl = DocxTemplate(_docx_with("{{ cycler.__init__.__globals__ }}"))
    with pytest.raises(SecurityError):
        tpl.render({}, jinja_env=_SANDBOX_ENV)


def test_docx_sandbox_allows_merge_field():
    """Een gewoon merge-veld rendert nog gewoon door de sandbox."""
    tpl = DocxTemplate(_docx_with("Beste {{ naam }}, uw dossier {{ nummer }}."))
    tpl.render({"naam": "Jan Jansen", "nummer": "2026-00001"}, jinja_env=_SANDBOX_ENV)
    out = io.BytesIO()
    tpl.save(out)
    out.seek(0)
    text = " ".join(p.text for p in Document(out).paragraphs)
    assert "Beste Jan Jansen" in text
    assert "2026-00001" in text


def test_html_env_blocks_ssti():
    """Het oude HTML-pad (service.jinja_env) is óók een sandbox."""
    with pytest.raises(SecurityError):
        jinja_env.from_string("{{ cycler.__init__.__globals__ }}").render()


def test_html_env_allows_normal_template():
    """Normale HTML-sjablonen + eigen filters blijven werken."""
    out = jinja_env.from_string("Bedrag: {{ x | currency }}").render(x=1234.5)
    assert "1.234,50" in out
