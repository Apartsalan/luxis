"""Generate .docx templates for the incasso document generation module.

Run once to create template files:
    python templates/_generate_templates.py

Uses docxtpl Jinja2 syntax:
  {{ variable }}           — simple substitution
  {%tr for x in list %}    — table row loop (start of row)
  {%tr endfor %}           — end of table row loop (own row, removed after render)
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

TEMPLATES_DIR = Path(__file__).parent


def _hdr(cell, text):
    """Style a table header cell."""
    p = cell.paragraphs[0]
    p.text = ""
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)


def _cel(cell, text, bold=False, right=False):
    """Style a table data cell."""
    p = cell.paragraphs[0]
    p.text = ""
    if right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(10)


def _p(doc, text, bold=False, size=Pt(11), color=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = size
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def _margins(doc):
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)


def _add_signature(doc, greeting="Hoogachtend,"):
    """Add Lisanne's full signature block — matches BaseNet style."""
    _p(doc, "")
    _p(doc, greeting)
    _p(doc, "")
    _p(doc, "Mevr. mr. L. Kesting", bold=True)
    _p(doc, "INCASSO ADVOCAAT | DEBT COLLECTION ATTORNEY",
       size=Pt(9), color=RGBColor(0x66, 0x66, 0x66))
    _p(doc, "")
    _p(doc, "{{ kantoor.adres }}", size=Pt(9))
    _p(doc, "{{ kantoor.postcode_stad }}", size=Pt(9))
    _p(doc, "E: {{ kantoor.email }}", size=Pt(9))
    _p(doc, "W: www.kestinglegal.nl", size=Pt(9))


def _add_schuldhulp(doc):
    """Wettelijk verplicht schuldhulpblok + disclaimer."""
    _p(doc, "")
    _p(doc,
       "Heeft u financi\u00eble zorgen en ziet u geen uitweg meer? "
       "Wij informeren u graag over uw rechten als schuldenaar: "
       "kestinglegal.nl/debiteuren. Voor schuldhulpverlening kunt u "
       "terecht bij uw gemeente. Heeft u dringend emotionele steun "
       "nodig? Bel dan gratis en anoniem met Stichting 113 "
       "Zelfmoordpreventie via 0800-0113 of kijk op www.113.nl.",
       size=Pt(8), color=RGBColor(0x88, 0x88, 0x88))
    _p(doc,
       "Disclaimer - De informatie verzonden met dit bericht is "
       "uitsluitend bestemd voor de geadresseerde(n) en kan "
       "persoonlijke of vertrouwelijke informatie bevatten, "
       "beschermd door een beroepsgeheim.",
       size=Pt(8), color=RGBColor(0xAA, 0xAA, 0xAA))


def _add_betaling_instructie(doc, termijn="2 DAGEN"):
    """Betalingsinstructie + IBAN derdengelden."""
    _p(doc,
       f"Hierbij sommeer ik u andermaal het bovengenoemd totaalbedrag "
       f"ad {{{{ totaal_openstaand }}}} UITERLIJK BINNEN {termijn} NA "
       "HEDEN te hebben bijgeschreven op de derdengeldenrekening van "
       "mijn kantoor IBAN: NL20 RABO 0388 5065 20 t.n.v. Stichting "
       "Beheer Derdengelden Kesting Legal onder vermelding van het "
       "kenmerk {{ zaak.zaaknummer }}.")


def _add_claims_table(doc):
    """Add a claims table with {%tr %} row loop.

    docxtpl replaces the ENTIRE <w:tr> containing {%tr} with just the Jinja2
    directive. So {%tr for %} and {%tr endfor %} must each be in their OWN
    row, separate from the data row.

    Structure:
      Row 0: Header (Omschrijving, Factuurnummer, Verzuimdatum, Hoofdsom)
      Row 1: {%tr for v in vorderingen %} — tag-only row, removed by docxtpl
      Row 2: Data row ({{ v.beschrijving }}, {{ v.factuurnummer }}, etc.)
      Row 3: {%tr endfor %} — tag-only row, removed by docxtpl
    """
    t = doc.add_table(rows=4, cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(["Omschrijving", "Factuurnummer", "Verzuimdatum", "Hoofdsom"]):
        _hdr(t.rows[0].cells[i], h)

    # For-loop opening tag (own row — gets replaced by {% for %} directive)
    t.rows[1].cells[0].text = "{%tr for v in vorderingen %}"

    # Data row (loop body — this row gets repeated for each item)
    r2 = t.rows[2]
    _cel(r2.cells[0], "{{ v.beschrijving }}")
    _cel(r2.cells[1], "{{ v.factuurnummer }}")
    _cel(r2.cells[2], "{{ v.verzuimdatum }}")
    _cel(r2.cells[3], "{{ v.hoofdsom }}", right=True)

    # End-loop tag (own row — gets replaced by {% endfor %} directive)
    t.rows[3].cells[0].text = "{%tr endfor %}"

    return t


def _add_summary_2col(doc, rows_data):
    """Add a 2-column summary table with label/value pairs."""
    t = doc.add_table(rows=len(rows_data), cols=2)
    t.style = "Table Grid"
    for i, (label, value, bold) in enumerate(rows_data):
        _cel(t.rows[i].cells[0], label, bold=bold)
        _cel(t.rows[i].cells[1], value, bold=bold, right=True)
    return t


# ═══════════════════════════════════════════════════════════════════════════════
# 14-DAGENBRIEF (art. 6:96 lid 6 BW)
# ═══════════════════════════════════════════════════════════════════════════════


def create_14_dagenbrief():
    doc = Document()
    _margins(doc)

    # Kantoor header
    _p(doc, "{{ kantoor.naam }}", bold=True)
    _p(doc, "{{ kantoor.adres }}")
    _p(doc, "{{ kantoor.postcode_stad }}")
    _p(doc, "Tel: {{ kantoor.telefoon }}")
    _p(doc, "")

    # Recipient
    _p(doc, "{{ wederpartij.naam }}", bold=True)
    _p(doc, "{{ wederpartij.adres }}")
    _p(doc, "{{ wederpartij.postcode_stad }}")
    _p(doc, "")

    _p(doc, "{{ vandaag }}")
    _p(doc, "")

    _p(doc, "Betreft: Ingebrekestelling en aanmaning tot betaling", bold=True)
    _p(doc, "Ons kenmerk: {{ zaak.zaaknummer }}")
    _p(doc, "{{ zaak.referentie_regel }}")
    _p(doc, "")

    _p(doc, "Geachte heer/mevrouw,")
    _p(doc, "")
    _p(doc,
        "Namens onze cli\u00ebnt, {{ client.naam }}, sommeer ik u hierbij tot betaling "
        "van het hieronder gespecificeerde bedrag. Ondanks eerdere verzoeken is betaling "
        "tot op heden uitgebleven.")

    doc.add_heading("Specificatie vordering(en)", level=2)
    _add_claims_table(doc)
    _p(doc, "")

    _add_summary_2col(doc, [
        ("Totaal hoofdsom", "{{ totaal_hoofdsom }}", True),
        ("Wettelijke rente t/m {{ vandaag }}", "{{ totaal_rente }}", False),
        ("Buitengerechtelijke incassokosten (BIK)", "{{ bik_bedrag }}", False),
        ("{{ btw_regel_label }}", "{{ btw_regel_bedrag }}", False),
        ("Totaal verschuldigd", "{{ totaal_verschuldigd }}", True),
    ])
    _p(doc, "")

    doc.add_heading("Wettelijke mededeling (art. 6:96 lid 6 BW)", level=2)
    _p(doc,
        "Hierbij wordt u gesommeerd het verschuldigde bedrag van "
        "{{ totaal_hoofdsom }} "
        "binnen veertien dagen na ontvangst van deze brief te voldoen "
        "op het rekeningnummer van onze cli\u00ebnt.")
    _p(doc,
        "Indien betaling binnen deze termijn uitblijft, bent u naast de hoofdsom "
        "tevens de buitengerechtelijke incassokosten verschuldigd ter hoogte van "
        "{{ bik_bedrag }}{{ btw_toelichting }}, "
        "alsmede de wettelijke rente vanaf de verzuimdatum.")
    _p(doc, "")
    _p(doc,
        "Ik verzoek u het totaal verschuldigde bedrag van "
        "{{ totaal_verschuldigd }} "
        "binnen veertien dagen na dagtekening van deze brief over te maken.")
    _p(doc,
        "Bij gebreke van tijdige betaling zullen wij zonder nadere aankondiging "
        "rechtsmaatregelen treffen, waarvan de kosten eveneens voor uw rekening komen.")
    _p(doc, "")
    _add_signature(doc)
    _add_schuldhulp(doc)

    doc.save(str(TEMPLATES_DIR / "14_dagenbrief.docx"))
    print("  14_dagenbrief.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# SOMMATIE
# ═══════════════════════════════════════════════════════════════════════════════


def create_sommatie():
    doc = Document()
    _margins(doc)

    _p(doc, "{{ kantoor.naam }}", bold=True)
    _p(doc, "{{ kantoor.adres }}")
    _p(doc, "{{ kantoor.postcode_stad }}")
    _p(doc, "Tel: {{ kantoor.telefoon }}")
    _p(doc, "")

    _p(doc, "AANGETEKEND", bold=True, size=Pt(12))
    _p(doc, "")
    _p(doc, "{{ wederpartij.naam }}", bold=True)
    _p(doc, "{{ wederpartij.adres }}")
    _p(doc, "{{ wederpartij.postcode_stad }}")
    _p(doc, "")

    _p(doc, "{{ vandaag }}")
    _p(doc, "")

    _p(doc, "Betreft: SOMMATIE \u2014 Laatste aanmaning voor dagvaarding", bold=True)
    _p(doc, "Ons kenmerk: {{ zaak.zaaknummer }}")
    _p(doc, "{{ zaak.referentie_regel }}")
    _p(doc, "")

    _p(doc, "Geachte heer/mevrouw,")
    _p(doc, "")
    _p(doc,
        "Tot mij heeft zich gewend {{ client.naam }} met het verzoek u te sommeren "
        "tot betaling van het hierna te noemen bedrag. Tot op heden is betaling, "
        "ondanks eerdere aanmaningen, uitgebleven.")

    doc.add_heading("Specificatie vordering", level=2)
    _add_claims_table(doc)
    _p(doc, "")

    _add_summary_2col(doc, [
        ("Hoofdsom", "{{ totaal_hoofdsom }}", True),
        ("Wettelijke rente t/m {{ vandaag }}", "{{ totaal_rente }}", False),
        ("Buitengerechtelijke incassokosten", "{{ bik_bedrag }}", False),
        ("{{ btw_regel_label }}", "{{ btw_regel_bedrag }}", False),
        ("{{ betalingen_regel_label }}", "{{ betalingen_regel_bedrag }}", False),
        ("Totaal thans verschuldigd", "{{ totaal_openstaand }}", True),
    ])
    _p(doc, "")

    doc.add_heading("SOMMATIE", level=2)
    _p(doc,
        "Ik sommeer u hierbij om het verschuldigde bedrag van "
        "{{ totaal_openstaand }} "
        "binnen acht (8) dagen na dagtekening van deze brief over te maken.")
    _p(doc,
        "Bij gebreke van tijdige en volledige betaling zal ik, zonder nadere "
        "aankondiging, overgaan tot het uitbrengen van een dagvaarding. "
        "De hieraan verbonden kosten (griffierecht, deurwaarderskosten, "
        "proceskosten) komen volledig voor uw rekening.")
    _p(doc, "")
    _p(doc,
        "Ik vertrouw erop u hiermee voldoende te hebben ge\u00efnformeerd en zie uw "
        "betaling met spoed tegemoet.")
    _p(doc, "")
    _add_signature(doc)
    _add_schuldhulp(doc)

    doc.save(str(TEMPLATES_DIR / "sommatie.docx"))
    print("  sommatie.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# RENTEOVERZICHT
# ═══════════════════════════════════════════════════════════════════════════════


def create_renteoverzicht():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)

    doc.add_heading("Renteberekening / Specificatie", level=1)

    # Meta
    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    for i, (l, v) in enumerate([
        ("Zaaknummer:", "{{ zaak.zaaknummer }}"),
        ("Cli\u00ebnt:", "{{ client.naam }}"),
        ("Wederpartij:", "{{ wederpartij.naam }}"),
        ("Type rente:", "{{ rente_type_label }}"),
        ("Berekend t/m:", "{{ vandaag }}"),
    ]):
        _cel(meta.rows[i].cells[0], l, bold=True)
        _cel(meta.rows[i].cells[1], v)
    _p(doc, "")

    # Claims
    doc.add_heading("Vorderingen", level=2)
    ct = _add_claims_table(doc)

    # Totaal row (static)
    tr = ct.add_row()
    _cel(tr.cells[0], "Totaal hoofdsom", bold=True)
    tr.cells[1].text = ""
    tr.cells[2].text = ""
    _cel(tr.cells[3], "{{ totaal_hoofdsom }}", bold=True, right=True)
    _p(doc, "")

    # Interest per period
    doc.add_heading("Renteberekening per periode", level=2)

    it = doc.add_table(rows=5, cols=6)
    it.style = "Table Grid"
    it.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Van", "Tot", "Dagen", "Tarief", "Hoofdsom", "Rente"]):
        _hdr(it.rows[0].cells[i], h)

    # For-loop opening tag (own row)
    it.rows[1].cells[0].text = "{%tr for r in rente_regels %}"

    # Data row (loop body)
    _cel(it.rows[2].cells[0], "{{ r.van }}")
    _cel(it.rows[2].cells[1], "{{ r.tot }}")
    _cel(it.rows[2].cells[2], "{{ r.dagen }}")
    _cel(it.rows[2].cells[3], "{{ r.tarief }}")
    _cel(it.rows[2].cells[4], "{{ r.hoofdsom }}")
    _cel(it.rows[2].cells[5], "{{ r.rente }}")

    # Endfor row
    it.rows[3].cells[0].text = "{%tr endfor %}"

    # Total
    _cel(it.rows[4].cells[0], "Totaal rente", bold=True)
    for c in it.rows[4].cells[1:5]:
        c.text = ""
    _cel(it.rows[4].cells[5], "{{ totaal_rente }}", bold=True, right=True)
    _p(doc, "")

    # BIK
    doc.add_heading("Buitengerechtelijke incassokosten (BIK)", level=2)
    _add_summary_2col(doc, [
        ("BIK (excl. BTW)", "{{ bik_bedrag }}", True),
        ("{{ btw_regel_label }}", "{{ btw_regel_bedrag }}", False),
        ("{{ bik_incl_label }}", "{{ bik_incl_bedrag }}", False),
    ])
    _p(doc, "")

    # Payments
    doc.add_heading("{{ betalingen_kop }}", level=2)
    pt = doc.add_table(rows=4, cols=3)
    pt.style = "Table Grid"
    for i, h in enumerate(["Datum", "Omschrijving", "Bedrag"]):
        _hdr(pt.rows[0].cells[i], h)

    # For-loop opening tag (own row)
    pt.rows[1].cells[0].text = "{%tr for b in betalingen %}"

    # Data row (loop body)
    _cel(pt.rows[2].cells[0], "{{ b.datum }}")
    _cel(pt.rows[2].cells[1], "{{ b.beschrijving }}")
    _cel(pt.rows[2].cells[2], "{{ b.bedrag }}")

    # Endfor row
    pt.rows[3].cells[0].text = "{%tr endfor %}"
    _p(doc, "")

    # Grand total
    doc.add_heading("Totaaloverzicht", level=2)
    _add_summary_2col(doc, [
        ("Hoofdsom", "{{ totaal_hoofdsom }}", False),
        ("Rente t/m {{ vandaag }}", "{{ totaal_rente }}", False),
        ("Buitengerechtelijke incassokosten", "{{ totaal_bik }}", False),
        ("Subtotaal", "{{ subtotaal }}", True),
        ("{{ betalingen_aftrek_label }}", "{{ betalingen_aftrek_bedrag }}", False),
        ("Totaal verschuldigd", "{{ totaal_openstaand }}", True),
    ])
    _p(doc, "")

    _p(doc, "")
    _p(doc, "Mevr. mr. L. Kesting", bold=True, size=Pt(9))
    _p(doc, "Kesting Legal | IJsbaanpad 9 | 1076 CV Amsterdam",
       size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))
    _p(doc,
       "Berekening gegenereerd op {{ vandaag }} | "
       "Zaaknummer: {{ zaak.zaaknummer }}",
       size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

    doc.save(str(TEMPLATES_DIR / "renteoverzicht.docx"))
    print("  renteoverzicht.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# HERINNERING
# ═══════════════════════════════════════════════════════════════════════════════


def create_herinnering():
    doc = Document()
    _margins(doc)

    # Kantoor header
    _p(doc, "{{ kantoor.naam }}", bold=True)
    _p(doc, "{{ kantoor.adres }}")
    _p(doc, "{{ kantoor.postcode_stad }}")
    _p(doc, "Tel: {{ kantoor.telefoon }}")
    _p(doc, "")

    # Recipient
    _p(doc, "{{ wederpartij.naam }}", bold=True)
    _p(doc, "{{ wederpartij.adres }}")
    _p(doc, "{{ wederpartij.postcode_stad }}")
    _p(doc, "")

    _p(doc, "{{ vandaag }}")
    _p(doc, "")

    _p(doc, "Betreft: Herinnering openstaande vordering", bold=True)
    _p(doc, "Ons kenmerk: {{ zaak.zaaknummer }}")
    _p(doc, "{{ zaak.referentie_regel }}")
    _p(doc, "")

    _p(doc, "Geachte heer/mevrouw,")
    _p(doc, "")
    _p(doc,
        "Wij constateren dat de onderstaande vordering(en) van onze "
        "cli\u00ebnt {{ client.naam }} nog niet zijn voldaan. "
        "Wij verzoeken u vriendelijk het openstaande bedrag "
        "binnen 14 dagen te voldoen.")
    _p(doc, "")

    doc.add_heading("Specificatie", level=2)
    _add_summary_2col(doc, [
        ("Hoofdsom", "{{ totaal_hoofdsom }}", True),
        ("Rente t/m {{ vandaag }}", "{{ totaal_rente }}", False),
        ("Totaal openstaand", "{{ totaal_openstaand }}", True),
    ])
    _p(doc, "")

    _p(doc,
        "Wij verzoeken u het bedrag van {{ totaal_openstaand }} "
        "voor {{ termijn_14_dagen }} over te maken op IBAN "
        "{{ kantoor.iban }} t.n.v. {{ kantoor.naam }}, onder vermelding "
        "van zaaknummer {{ zaak.zaaknummer }}.")
    _p(doc, "")
    _add_signature(doc, "Met vriendelijke groet,")
    _add_schuldhulp(doc)

    doc.save(str(TEMPLATES_DIR / "herinnering.docx"))
    print("  herinnering.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# AANMANING
# ═══════════════════════════════════════════════════════════════════════════════


def create_aanmaning():
    doc = Document()
    _margins(doc)

    _p(doc, "{{ kantoor.naam }}", bold=True)
    _p(doc, "{{ kantoor.adres }}")
    _p(doc, "{{ kantoor.postcode_stad }}")
    _p(doc, "Tel: {{ kantoor.telefoon }}")
    _p(doc, "")

    _p(doc, "{{ wederpartij.naam }}", bold=True)
    _p(doc, "{{ wederpartij.adres }}")
    _p(doc, "{{ wederpartij.postcode_stad }}")
    _p(doc, "")

    _p(doc, "{{ vandaag }}")
    _p(doc, "")

    _p(doc, "Betreft: Aanmaning tot betaling", bold=True)
    _p(doc, "Ons kenmerk: {{ zaak.zaaknummer }}")
    _p(doc, "{{ zaak.referentie_regel }}")
    _p(doc, "")

    _p(doc, "Geachte heer/mevrouw,")
    _p(doc, "")
    _p(doc,
        "Ondanks onze eerdere herinnering constateren wij dat de "
        "vordering(en) van onze cli\u00ebnt {{ client.naam }} nog steeds "
        "niet zijn voldaan.")
    _p(doc, "")

    doc.add_heading("Specificatie", level=2)
    _add_claims_table(doc)
    _p(doc, "")

    _add_summary_2col(doc, [
        ("Hoofdsom", "{{ totaal_hoofdsom }}", True),
        ("Rente t/m {{ vandaag }}", "{{ totaal_rente }}", False),
        ("BIK", "{{ bik_bedrag }}", False),
        ("{{ btw_regel_label }}", "{{ btw_regel_bedrag }}", False),
        ("Totaal verschuldigd", "{{ totaal_verschuldigd }}", True),
        ("{{ betalingen_aftrek_label }}", "{{ betalingen_aftrek_bedrag }}", False),
        ("Totaal openstaand", "{{ totaal_openstaand }}", True),
    ])
    _p(doc, "")

    _p(doc,
        "Wij sommeren u het bedrag van {{ totaal_openstaand }} uiterlijk "
        "{{ termijn_14_dagen }} te voldoen op IBAN {{ kantoor.iban }} "
        "t.n.v. {{ kantoor.naam }}, onder vermelding van zaaknummer "
        "{{ zaak.zaaknummer }}.")
    _p(doc, "")
    _p(doc,
        "Bij gebreke van tijdige betaling zullen wij zonder nadere "
        "aankondiging rechtsmaatregelen treffen.")
    _p(doc, "")
    _add_signature(doc)
    _add_schuldhulp(doc)

    doc.save(str(TEMPLATES_DIR / "aanmaning.docx"))
    print("  aanmaning.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# TWEEDE SOMMATIE
# ═══════════════════════════════════════════════════════════════════════════════


def create_tweede_sommatie():
    doc = Document()
    _margins(doc)

    _p(doc, "{{ kantoor.naam }}", bold=True)
    _p(doc, "{{ kantoor.adres }}")
    _p(doc, "{{ kantoor.postcode_stad }}")
    _p(doc, "Tel: {{ kantoor.telefoon }}")
    _p(doc, "")

    _p(doc, "AANGETEKEND", bold=True, size=Pt(12))
    _p(doc, "")
    _p(doc, "{{ wederpartij.naam }}", bold=True)
    _p(doc, "{{ wederpartij.adres }}")
    _p(doc, "{{ wederpartij.postcode_stad }}")
    _p(doc, "")

    _p(doc, "{{ vandaag }}")
    _p(doc, "")

    _p(doc, "Betreft: TWEEDE SOMMATIE \u2014 Laatste gelegenheid", bold=True)
    _p(doc, "Ons kenmerk: {{ zaak.zaaknummer }}")
    _p(doc, "{{ zaak.referentie_regel }}")
    _p(doc, "")

    _p(doc, "Geachte heer/mevrouw,")
    _p(doc, "")
    _p(doc,
        "Ondanks onze eerdere sommatie heeft u nog steeds niet voldaan "
        "aan de vordering(en) van onze cli\u00ebnt {{ client.naam }}. "
        "Dit is uw laatste gelegenheid om vrijwillig te betalen.")
    _p(doc, "")

    doc.add_heading("Specificatie", level=2)
    _add_summary_2col(doc, [
        ("Hoofdsom", "{{ totaal_hoofdsom }}", True),
        ("Rente t/m {{ vandaag }}", "{{ totaal_rente }}", False),
        ("BIK", "{{ bik_bedrag }}", False),
        ("{{ btw_regel_label }}", "{{ btw_regel_bedrag }}", False),
        ("Totaal verschuldigd", "{{ totaal_verschuldigd }}", True),
        ("{{ betalingen_aftrek_label }}", "{{ betalingen_aftrek_bedrag }}", False),
        ("Totaal openstaand", "{{ totaal_openstaand }}", True),
    ])
    _p(doc, "")

    _p(doc,
        "Wij sommeren u \u2014 voor de laatste maal \u2014 het bedrag van "
        "{{ totaal_openstaand }} uiterlijk {{ termijn_14_dagen }} te voldoen "
        "op IBAN {{ kantoor.iban }} t.n.v. {{ kantoor.naam }}, onder vermelding "
        "van zaaknummer {{ zaak.zaaknummer }}.")
    _p(doc, "")
    _p(doc,
        "Bij gebreke van tijdige betaling zullen wij zonder nadere "
        "aankondiging tot dagvaarding overgaan. De kosten hiervan "
        "(griffierecht, deurwaarderskosten, proceskosten) komen "
        "volledig voor uw rekening.")
    _p(doc, "")
    _add_signature(doc)
    _add_schuldhulp(doc)

    doc.save(str(TEMPLATES_DIR / "tweede_sommatie.docx"))
    print("  tweede_sommatie.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# DAGVAARDING (concept)
# ═══════════════════════════════════════════════════════════════════════════════


def create_dagvaarding():
    doc = Document()
    _margins(doc)

    doc.add_heading("CONCEPT DAGVAARDING", level=1)
    _p(doc, "")

    _p(doc,
        "Heden, {{ vandaag }}, ten verzoeke van {{ client.naam }}, "
        "gevestigd/wonende te {{ client.postcode_stad }}, {{ client.adres }}, "
        "te dezer zake domicilie kiezende ten kantore van {{ kantoor.naam }}, "
        "{{ kantoor.adres }}, {{ kantoor.postcode_stad }},")
    _p(doc, "")
    _p(doc,
        "heb ik, [NAAM DEURWAARDER], als gerechtsdeurwaarder, "
        "ge\u00ebxploiteerd aan:")
    _p(doc, "")
    _p(doc,
        "{{ wederpartij.naam }}, gevestigd/wonende te "
        "{{ wederpartij.postcode_stad }}, {{ wederpartij.adres }},")
    _p(doc, "")
    _p(doc, "hierna te noemen: \"gedaagde\",")
    _p(doc, "")

    doc.add_heading("om", level=2)
    _p(doc,
        "op [DATUM ZITTING] om [TIJDSTIP] uur, niet in persoon doch "
        "vertegenwoordigd door een advocaat, te verschijnen ter "
        "terechtzitting van de Rechtbank [RECHTBANK], sector [SECTOR].")
    _p(doc, "")

    doc.add_heading("TENEINDE", level=2)
    _p(doc, "")

    doc.add_heading("I. Feiten", level=2)
    _p(doc,
        "Eiser heeft vordering(en) op gedaagde uit hoofde van: "
        "{{ zaak.omschrijving }}")
    _p(doc, "")

    doc.add_heading("Vorderingen", level=3)
    _add_claims_table(doc)
    _p(doc, "")

    doc.add_heading("II. Rente", level=2)
    _p(doc,
        "Over de hoofdsom is gedaagde {{ rente_type_label }} verschuldigd "
        "ten bedrage van {{ totaal_rente }}.")
    _p(doc, "")

    doc.add_heading("III. Buitengerechtelijke kosten", level=2)
    _p(doc,
        "Eiser maakt aanspraak op buitengerechtelijke incassokosten "
        "conform de staffel BIK (art. 6:96 BW) ten bedrage van "
        "{{ bik_bedrag }}{{ btw_toelichting }}.")
    _p(doc, "")

    doc.add_heading("IV. Specificatie", level=2)
    _add_summary_2col(doc, [
        ("Hoofdsom", "{{ totaal_hoofdsom }}", True),
        ("Rente t/m {{ vandaag }}", "{{ totaal_rente }}", False),
        ("BIK", "{{ bik_bedrag }}", False),
        ("{{ btw_regel_label }}", "{{ btw_regel_bedrag }}", False),
        ("Totaal verschuldigd", "{{ totaal_verschuldigd }}", True),
        ("{{ betalingen_aftrek_label }}", "{{ betalingen_aftrek_bedrag }}", False),
        ("Totaal openstaand", "{{ totaal_openstaand }}", True),
    ])
    _p(doc, "")

    doc.add_heading("MITSDIEN", level=2)
    _p(doc, "het de Rechtbank behage, bij vonnis uitvoerbaar bij voorraad:")
    _p(doc,
        "gedaagde te veroordelen om aan eiser te betalen het bedrag van "
        "{{ totaal_openstaand }}, te vermeerderen met de {{ rente_type_label }} "
        "over de hoofdsom van {{ totaal_hoofdsom }} vanaf heden tot aan de "
        "dag der algehele voldoening,")
    _p(doc, "gedaagde te veroordelen in de kosten van dit geding.")
    _p(doc, "")

    _p(doc, "De kosten van dit exploot zijn: [KOSTEN]",
       size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))
    _p(doc, "")
    _p(doc, "Mevr. mr. L. Kesting", bold=True)
    _p(doc, "Kesting Legal")
    _p(doc, "Advocaat")

    doc.save(str(TEMPLATES_DIR / "dagvaarding.docx"))
    print("  dagvaarding.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# VERZOEKSCHRIFT FAILLISSEMENT (concept — meegestuurd als bijlage bij dreigbrief)
# ═══════════════════════════════════════════════════════════════════════════════


def create_verzoekschrift_faillissement():
    """Generate concept verzoekschrift tot faillietverklaring.

    Drie secties (op aparte pagina's):
      1. Begeleidende brief (persoonlijk & vertrouwelijk, sommatie 3 dagen)
      2. Verzoekschrift ex art. 1 Fw (legal petition)
      3. Slotpagina met laatste waarschuwing + contact
    """
    doc = Document()
    _margins(doc)

    # ─── Sectie 1: Begeleidende brief ──────────────────────────────────────
    _p(doc, "PERSOONLIJK & VERTROUWELIJK", bold=True, size=Pt(10))
    _p(doc, "")

    # Kantoor + wederpartij header
    _p(doc, "{{ kantoor.naam }}", bold=True)
    _p(doc, "{{ kantoor.adres }}")
    _p(doc, "{{ kantoor.postcode_stad }}")
    _p(doc, "")
    _p(doc, "{{ wederpartij.naam }}", bold=True)
    _p(doc, "{{ wederpartij.adres }}")
    _p(doc, "{{ wederpartij.postcode_stad }}")
    _p(doc, "")
    _p(doc, "Datum: {{ vandaag }}")
    _p(doc, "Onze referentie: {{ zaak.zaaknummer }}")
    _p(doc, "Inzake: Verzoekschrift faillissement")
    _p(doc, "")

    _p(doc, "Geachte heer, mevrouw {{ wederpartij.naam }},")
    _p(doc, "")
    _p(doc,
        "{{ client.naam }}, gevestigd te {{ client.postcode_stad }} "
        "(hierna te noemen: 'cli\u00ebnte') heeft zich tot mij gewend "
        "aangezien de hieronder vermelde vordering ten hoogte van "
        "{{ totaal_openstaand }} nog niet door u is voldaan.")
    _p(doc, "")
    _p(doc,
        "Ik stel u hierbij op de hoogte dat de vordering wordt gestuit "
        "conform artikel 3:317 BW. Het recht op nakoming wordt hiermee "
        "ondubbelzinnig voorbehouden.")
    _p(doc, "")
    _p(doc,
        "Bijgaand bij deze brief treft u de bijbehorende facturen aan. "
        "Onderstaand zend ik u een specificatie van uw schuld:")
    _p(doc, "")

    _add_claims_table(doc)
    _p(doc, "")

    _add_summary_2col(doc, [
        ("Hoofdsom", "{{ totaal_hoofdsom }}", True),
        ("Rente t/m {{ vandaag }}", "{{ totaal_rente }}", False),
        ("Hoofdsom + rente", "{{ subtotaal }}", False),
        ("Incassokosten (BIK)", "{{ bik_bedrag }}", False),
        ("{{ btw_regel_label }}", "{{ btw_regel_bedrag }}", False),
        ("Totaal verschuldigd", "{{ totaal_verschuldigd }}", True),
        ("{{ betalingen_aftrek_label }}", "{{ betalingen_aftrek_bedrag }}", False),
        ("Te voldoen", "{{ totaal_openstaand }}", True),
    ])
    _p(doc, "")

    _p(doc,
        "Indien het voormelde bedrag niet binnen een termijn van "
        "DRIE DAGEN NA HEDEN NA BETEKENING VOOR 12:00 UUR is bijgeschreven "
        "op de derdengeldenrekening van mijn kantoor: IBAN NL20 RABO "
        "0388 5065 20 ten name van Stichting Beheer Derdengelden Kesting "
        "Legal onder vermelding van het kenmerk {{ zaak.zaaknummer }}, "
        "dan zal ik cli\u00ebnte adviseren het bij deze brief aangehechte "
        "concept verzoek tot faillietverklaring in te dienen.")
    _p(doc, "")
    _p(doc,
        "Lagere betalingen dan gevorderd zijn, worden altijd in mindering "
        "gebracht op eerst de verschenen en gevorderde buitengerechtelijke "
        "kosten, vervolgens op verschenen rente en als laatste op de "
        "hoofdsom als gevorderd.")
    _p(doc, "")
    _p(doc, "Kosten van de faillissementsaanvraag bedragen EUR 2.195,00.")
    _p(doc, "")
    _p(doc,
        "Voorts wijs ik u erop op voorhand op dat, indien de behandeling "
        "van de faillissementsaanvraag \u00e9\u00e9n of meerdere malen "
        "dient te worden aangehouden, \u00e9\u00e9n punt liquidatietarief "
        "is verschuldigd van \u20ac 412,61 inclusief BTW boven de kosten "
        "der aanvraag.")
    _p(doc, "")
    _p(doc, "In afwachting van uw spoedige betaling.")
    _p(doc, "")

    _add_signature(doc)

    # ─── Sectie 2: Verzoekschrift ex art. 1 Fw ────────────────────────────
    doc.add_page_break()

    doc.add_heading("Verzoekschrift tot faillietverklaring (ex. art. 1 Fw)",
                    level=1)
    _p(doc, "")
    _p(doc, "Aan de rechtbank [nader in te vullen]", bold=True)
    _p(doc, "")
    _p(doc, "Geeft eerbiedig te kennen:", bold=True)
    _p(doc, "")
    _p(doc,
        "De besloten vennootschap met beperkte aansprakelijkheid "
        "{{ client.naam }}, gevestigd te {{ client.postcode_stad }}, "
        "in deze zaak woonplaats kiezende te Amsterdam, ten kantore "
        "van {{ kantoor.naam }}, van wie tot advocaat wordt gesteld "
        "Mevr. mr. L. Kesting met het recht van substitutie.")
    _p(doc, "")
    _p(doc, "(\u201cverzoeker\u201d)", bold=True)
    _p(doc, "")
    _p(doc, "Dit verzoek is gericht tegen:", bold=True)
    _p(doc, "")
    _p(doc,
        "{{ wederpartij.naam }}, woonachtig en/of zaakdoende te "
        "{{ wederpartij.postcode_stad }}, aan {{ wederpartij.adres }} "
        "(bijlage 1).")
    _p(doc, "")
    _p(doc, "(\u201cverweerder\u201d)", bold=True)
    _p(doc, "")
    _p(doc,
        "Dit rekest strekt tot faillietverklaring van verweerder, "
        "waartoe verzoeker het navolgende aanvoert:")
    _p(doc, "")

    _p(doc,
        "1. Verzoeker heeft een opeisbare vordering uit hoofde van "
        "geleverde zaken en/of diensten en daarbij horende facturen "
        "(bijlage 2).")
    _p(doc,
        "2. Verweerder blijft ondanks herhaaldelijke sommaties in "
        "gebreke de verschuldigde bedragen aan verzoeker te voldoen.")
    _p(doc,
        "3. De opeisbare vordering van Verzoeker op Verweerder "
        "bedraagt thans:")
    _p(doc, "")

    _add_claims_table(doc)
    _p(doc, "")

    _add_summary_2col(doc, [
        ("Hoofdsom", "{{ totaal_hoofdsom }}", True),
        ("Rente", "{{ totaal_rente }}", False),
        ("Hoofdsom + rente", "{{ subtotaal }}", False),
        ("Incassokosten", "{{ bik_bedrag }}", False),
        ("{{ btw_regel_label }}", "{{ btw_regel_bedrag }}", False),
        ("Totaal", "{{ totaal_verschuldigd }}", True),
        ("Te voldoen", "{{ totaal_openstaand }}", True),
    ])
    _p(doc, "")

    _p(doc,
        "4. Verweerder laat daarnaast nog \u00e9\u00e9n of meerdere "
        "vorderingen van \u00e9\u00e9n of meer andere schuldeisers "
        "onbetaald.")
    _p(doc,
        "5. Verweerder verkeert dus in de toestand dat hij heeft "
        "opgehouden te betalen. Verzoeker is derhalve gerechtigd de "
        "faillietverklaring van Verweerder te verzoeken.")
    _p(doc,
        "6. De kosten van de faillissementsaanvraag bedragen "
        "\u20ac 2.195,00.")
    _p(doc,
        "7. De voornaamste belangen van verweerder zijn gelegen in "
        "Nederland. Verweerder is woonachtig/zaakdoende/gevestigd te "
        "{{ wederpartij.postcode_stad }}, zodat Uw Rechtbank bevoegd is.")
    _p(doc, "")

    _p(doc, "Redenen waarom:", bold=True)
    _p(doc,
        "Verzoeker uw Rechtbank verzoekt verweerder in staat van "
        "faillissement te verklaren.")
    _p(doc, "")
    _p(doc, "Mevr. mr. L. Kesting", bold=True)
    _p(doc, "{{ kantoor.naam }}")
    _p(doc, "Advocaat")

    # ─── Sectie 3: Slotpagina ─────────────────────────────────────────────
    doc.add_page_break()

    _p(doc,
       "U kunt de faillissementszitting nog voorkomen door uiterlijk "
       "3 dagen voor de door de rechtbank aangekondigde zittingsdatum "
       "het bedrag van {{ totaal_openstaand }} bij te laten schrijven "
       "op de derdengeldenrekening van mijn kantoor IBAN NL20 RABO 0388 "
       "5065 20 t.n.v. Stichting Beheer Derdengelden Kesting Legal "
       "onder vermelding van het kenmerk {{ zaak.zaaknummer }}.",
       bold=True)
    _p(doc, "")
    _p(doc,
        "Een deelbetaling zal eerst in mindering strekken op de kosten "
        "en rente.")
    _p(doc, "")
    _p(doc,
        "Dit dossier is in behandeling bij mevrouw mr. L. Kesting van "
        "{{ kantoor.naam }} te {{ kantoor.postcode_stad }} aan "
        "{{ kantoor.adres }}.")
    _p(doc, "")
    _p(doc, "E: {{ kantoor.email }} | T: {{ kantoor.telefoon }}",
       size=Pt(9))

    doc.save(str(TEMPLATES_DIR / "verzoekschrift_faillissement.docx"))
    print("  verzoekschrift_faillissement.docx")


if __name__ == "__main__":
    print("Generating .docx templates...")
    create_herinnering()
    create_aanmaning()
    create_14_dagenbrief()
    create_sommatie()
    create_tweede_sommatie()
    create_dagvaarding()
    create_renteoverzicht()
    create_verzoekschrift_faillissement()
    print("Done!")
